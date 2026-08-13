from __future__ import annotations

from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SCREENSHOTS = ROOT / "output" / "playwright" / "review-final"
OUTPUT_DIR = ROOT / "docs" / "deliverables"
OUTPUT = OUTPUT_DIR / "xinhowsaas-system-review-and-operation-manual-2026-08-13-v2.docx"

NAVY = "08263A"
NAVY_2 = "12395A"
BLUE = "2178C7"
PALE_BLUE = "EAF3FB"
LIGHT_BLUE = "F4F8FC"
SLATE = "475569"
PALE_SLATE = "F1F5F9"
GREEN = "047857"
PALE_GREEN = "ECFDF5"
AMBER = "B45309"
PALE_AMBER = "FFFBEB"
RED = "B91C1C"
PALE_RED = "FEF2F2"
WHITE = "FFFFFF"
BLACK = "0F172A"


def set_font(run, name: str = "Microsoft JhengHei", size: float | None = None, bold: bool | None = None, color: str | None = None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    return run


def shade_cell(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top: int = 90, start: int = 110, bottom: int = 90, end: int = 110):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:tblHeader")) is None:
        tbl_header = OxmlElement("w:tblHeader")
        tbl_header.set(qn("w:val"), "true")
        tr_pr.append(tbl_header)


def finalize_table_accessibility(doc: Document):
    """Mark the first row of every layout/data table for assistive technology."""
    for table in doc.tables:
        if table.rows:
            set_repeat_table_header(table.rows[0])

    for section in doc.sections:
        for container in (section.header, section.footer):
            for table in container.tables:
                if table.rows:
                    set_repeat_table_header(table.rows[0])


def set_paragraph_keep(paragraph, keep_with_next: bool = True, keep_together: bool = True):
    p_pr = paragraph._p.get_or_add_pPr()
    if keep_with_next:
        p_pr.append(OxmlElement("w:keepNext"))
    if keep_together:
        p_pr.append(OxmlElement("w:keepLines"))


def add_field(paragraph, instruction: str):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for item in (begin, instr, separate, text, end):
        run._r.append(item)


def configure_styles(doc: Document):
    normal = doc.styles["Normal"]
    normal.font.name = "Microsoft JhengHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(SLATE)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.22

    configs = {
        "Title": (30, NAVY, 1, 12),
        "Heading 1": (20, NAVY, 16, 9),
        "Heading 2": (15, BLUE, 12, 7),
        "Heading 3": (12, NAVY_2, 9, 5),
        "Caption": (9, SLATE, 4, 8),
    }
    for style_name, (size, color, before, after) in configs.items():
        style = doc.styles[style_name]
        style.font.name = "Microsoft JhengHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
        style.font.size = Pt(size)
        style.font.bold = style_name != "Caption"
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def configure_sections(doc: Document):
    section = doc.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.68)
    section.right_margin = Inches(0.68)
    section.header_distance = Inches(0.25)
    section.footer_distance = Inches(0.28)

    header = section.header
    table = header.add_table(rows=1, cols=2, width=Inches(6.95))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.columns[0].width = Inches(4.7)
    table.columns[1].width = Inches(2.25)
    left = table.cell(0, 0)
    right = table.cell(0, 1)
    left_p = left.paragraphs[0]
    right_p = right.paragraphs[0]
    set_font(left_p.add_run("XINHOW  ·  OPERATIONS GUIDE"), size=8.5, bold=True, color=NAVY)
    set_font(right_p.add_run("STAGING VERIFIED  |  2026-08-13"), size=7.5, bold=True, color=BLUE)
    right_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    footer = section.footer
    footer_table = footer.add_table(rows=1, cols=2, width=Inches(6.95))
    footer_table.columns[0].width = Inches(5.7)
    footer_table.columns[1].width = Inches(1.25)
    p_left = footer_table.cell(0, 0).paragraphs[0]
    p_right = footer_table.cell(0, 1).paragraphs[0]
    set_font(p_left.add_run("多品牌預約與報名 SaaS · 圖文操作與產品優化手冊"), size=8, color=SLATE)
    set_font(p_right.add_run("第 "), size=8, color=SLATE)
    add_field(p_right, "PAGE")
    set_font(p_right.add_run(" 頁"), size=8, color=SLATE)
    p_right.alignment = WD_ALIGN_PARAGRAPH.RIGHT


def add_cover(doc: Document):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    shade_cell(cell, NAVY)
    set_cell_margins(cell, top=300, start=320, bottom=300, end=320)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_font(p.add_run("XINHOW"), size=12, bold=True, color="7CC3FF")
    p = cell.add_paragraph()
    set_font(p.add_run("多品牌預約與報名 SaaS"), size=24, bold=True, color=WHITE)
    p = cell.add_paragraph()
    set_font(p.add_run("系統 Review、UI/UX 優化、角色流程與圖文操作手冊"), size=16, bold=True, color=WHITE)
    p = cell.add_paragraph()
    set_font(p.add_run("SYSTEM REVIEW  ·  ROLE PLAYBOOK  ·  OPERATIONS MANUAL"), size=8.5, bold=True, color="9DD3FF")

    doc.add_paragraph()
    meta = doc.add_table(rows=4, cols=2)
    meta.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta.style = "Table Grid"
    rows = [
        ("文件版本", "v2.0 · 2026-08-13"),
        ("驗證環境", "Railway staging · deployment f58f03c4-3758-4c2f-84ce-37f1f0513bf6"),
        ("適用身分", "系統管理者、系統員工、品牌管理者、品牌員工"),
        ("規格基準", "AGENTS.md 與 clinic-booking-spec-v3.md"),
    ]
    for row, (label, value) in zip(meta.rows, rows):
        shade_cell(row.cells[0], PALE_BLUE)
        for cell in row.cells:
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_font(row.cells[0].paragraphs[0].add_run(label), size=9, bold=True, color=NAVY)
        set_font(row.cells[1].paragraphs[0].add_run(value), size=9, color=SLATE)

    doc.add_paragraph()
    add_callout(
        doc,
        "文件定位",
        "本手冊的畫面全部取自 2026-08-13 最新 staging 實際介面。沒有使用 SVG、線框稿或虛構 UI。測試帳號只用於驗收，已從文件排除密碼。",
        "info",
    )
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_font(p.add_run("XINHOW PRODUCT & ENGINEERING REVIEW"), size=9, bold=True, color=BLUE)
    doc.add_page_break()


def add_callout(doc: Document, title: str, body: str, tone: str = "info"):
    fills = {"info": PALE_BLUE, "success": PALE_GREEN, "warning": PALE_AMBER, "danger": PALE_RED}
    colors = {"info": BLUE, "success": GREEN, "warning": AMBER, "danger": RED}
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    shade_cell(cell, fills[tone])
    set_cell_margins(cell, top=130, start=160, bottom=130, end=160)
    p = cell.paragraphs[0]
    set_font(p.add_run(title), size=10, bold=True, color=colors[tone])
    p = cell.add_paragraph()
    set_font(p.add_run(body), size=9.5, color=BLACK)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_bullets(doc: Document, items: Iterable[str], level: int = 0):
    for item in items:
        p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
        set_font(p.add_run(item), size=10, color=SLATE)
        p.paragraph_format.space_after = Pt(3)


def add_numbered(doc: Document, items: Iterable[str]):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        set_font(p.add_run(item), size=10, color=SLATE)
        p.paragraph_format.space_after = Pt(4)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_repeat_table_header(table.rows[0])
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        shade_cell(cell, NAVY)
        set_cell_margins(cell)
        set_font(cell.paragraphs[0].add_run(header), size=8.8, bold=True, color=WHITE)
    for values in rows:
        row = table.add_row()
        for i, value in enumerate(values):
            cell = row.cells[i]
            set_cell_margins(cell)
            if len(table.rows) % 2 == 1:
                shade_cell(cell, LIGHT_BLUE)
            set_font(cell.paragraphs[0].add_run(value), size=8.6, color=SLATE)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    if widths:
        for row in table.rows:
            for cell, width in zip(row.cells, widths):
                cell.width = Inches(width)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def add_figure(doc: Document, filename: str, caption: str, alt: str, width: float = 6.75):
    path = SCREENSHOTS / filename
    if not path.exists():
        raise FileNotFoundError(path)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_keep(p, keep_with_next=True, keep_together=True)
    run = p.add_run()
    shape = run.add_picture(str(path), width=Inches(width))
    shape._inline.docPr.set("descr", alt)
    cap = doc.add_paragraph(style="Caption")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(cap.add_run(caption), size=8.8, color=SLATE)
    return shape


def add_manual_screen(doc: Document, heading: str, filename: str, caption: str, purpose: str, steps: list[str], note: str | None = None):
    doc.add_heading(heading, level=2)
    p = doc.add_paragraph()
    set_font(p.add_run(purpose), size=10, color=SLATE)
    add_figure(doc, filename, caption, f"{heading}的 staging 實際畫面")
    doc.add_heading("操作步驟", level=3)
    add_numbered(doc, steps)
    if note:
        add_callout(doc, "操作提醒", note, "warning")


def add_section_break(doc: Document, title: str, subtitle: str):
    doc.add_page_break()
    p = doc.add_paragraph()
    set_font(p.add_run("XINHOW PLAYBOOK"), size=8, bold=True, color=BLUE)
    doc.add_heading(title, level=1)
    p = doc.add_paragraph()
    set_font(p.add_run(subtitle), size=11, color=SLATE)


def add_mobile_pair(doc: Document):
    doc.add_heading("手機版操作檢查", level=2)
    p = doc.add_paragraph()
    set_font(p.add_run("品牌工作台與設定分類在 390 × 844 px staging 視窗實測；手機以漢堡選單收合主導覽，任務卡與設定分類改為單欄。"), size=10, color=SLATE)
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for idx, (filename, alt) in enumerate([
        ("22-mobile-brand-dashboard.png", "手機版品牌今日工作台"),
        ("23-mobile-brand-settings.png", "手機版品牌設定分類"),
    ]):
        cell = table.cell(0, idx)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        shape = p.add_run().add_picture(str(SCREENSHOTS / filename), width=Inches(2.7))
        shape._inline.docPr.set("descr", alt)
    cap = doc.add_paragraph(style="Caption")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(cap.add_run("圖 21–22　品牌工作台與設定中心的手機版實際畫面"), size=8.8, color=SLATE)
    add_bullets(doc, [
        "手機先看今日任務與上線阻擋，不需要先展開完整側欄。",
        "設定分類維持清楚的工作區卡片；進入後才顯示該領域表單。",
        "正式驗收仍應在 iOS Safari 與 Android Chrome 各跑一次登入、預約、Rich Menu LIFF 與改期。",
    ])


def add_four_role_mobile_grid(doc: Document):
    doc.add_heading("四身分手機驗收", level=2)
    p = doc.add_paragraph()
    set_font(p.add_run("四個身分均以 390 × 844 px staging 視窗完成登入、落點、導覽與拒絕測試；所有畫面無水平溢出與 application error。"), size=10, color=SLATE)
    table = doc.add_table(rows=2, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    items = [
        ("27-mobile-system-reports.png", "系統管理者手機跨品牌報表"),
        ("35-mobile-csv-import.png", "品牌管理者手機 CSV 匯入"),
        ("38-mobile-system-employee.png", "系統員工手機權限拒絕提示"),
        ("40-mobile-brand-employee.png", "品牌員工手機權限拒絕提示"),
    ]
    for cell, (filename, alt) in zip((cell for row in table.rows for cell in row.cells), items):
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        shape = p.add_run().add_picture(str(SCREENSHOTS / filename), width=Inches(2.55))
        shape._inline.docPr.set("descr", alt)
        caption = cell.add_paragraph()
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_font(caption.add_run(alt), size=8.2, color=SLATE)
    add_callout(doc, "判定", "系統員工與品牌員工直接輸入未授權網址時，均由 server guard 返回正確工作區並顯示申請權限方式；不是單純把選單藏起來。", "success")


def add_toc(doc: Document):
    doc.add_heading("閱讀導覽", level=1)
    rows = [
        ["A", "先看結論", "系統 Review、角色模型、已完成優化與上線條件"],
        ["B", "系統管理者手冊", "登入、總覽、建立品牌、交付、系統人員權限"],
        ["C", "品牌管理者手冊", "啟用中心、設定、服務、排班、團隊權限"],
        ["D", "日常營運手冊", "預約、活動、報名、CRM Lite、報表"],
        ["E", "LINE 與 Rich Menu", "渠道驗證、草稿、圖片、發布、Alias、排程、回復"],
        ["F", "員工與手機操作", "最小權限、拒絕提示、行動版流程"],
        ["G", "產品路線", "競品對照、功能優先序、KPI 與暫不納入項目"],
        ["H", "驗收與故障排除", "staging 證據、正式外部渠道檢查、常見問題"],
    ]
    add_table(doc, ["章", "主題", "你會完成什麼"], rows, [0.45, 1.55, 4.7])
    add_callout(doc, "最快使用方式", "第一次導入先讀 A、B、C、E；日常櫃檯只需讀 D、F；工程與產品負責人再讀 G、H。", "info")
    doc.add_heading("三層工作邊界", level=2)
    add_table(doc, ["層級", "負責內容", "不應做的事"], [
        ["系統層", "品牌租戶、平台人員、營運健康、跨品牌報表、稽核與政策", "代替品牌處理顧客、預約或訊息內容"],
        ["品牌管理層", "服務、排班、入口、LINE、金流、員工權限與營運規則", "查看其他品牌資料或管理系統人員"],
        ["品牌營運層", "依員工授權處理預約、報名、顧客、客服與報表", "自行提升權限或修改品牌安全設定"],
    ], [1.1, 3.1, 2.5])


def build_document() -> Document:
    doc = Document()
    configure_styles(doc)
    configure_sections(doc)
    add_cover(doc)
    add_toc(doc)

    add_section_break(doc, "A. 先看結論", "這次不是增加更多選單，而是把系統改為任務導向，並保留既有 v3 功能與安全契約。")
    doc.add_heading("Review 結論", level=2)
    add_callout(doc, "整體判定：可進入受控試營運", "前後端核心、RLS、角色權限、預約競爭、報名與付款生命週期、通知去重及瀏覽器身分已在 staging 通過。正式 LINE、Email、金流與 DNS 仍需用品牌真實帳號完成外部 smoke test。", "success")
    add_table(doc, ["面向", "判定", "重點"], [
        ["租戶與個資", "通過", "anon 無法讀顧客 PII；跨品牌讀寫與偽造入口被拒絕"],
        ["角色權限", "通過", "兩個管理身分；員工採細權限；前端隱藏與 server guard 同時存在"],
        ["預約／報名", "通過", "時間制、場次制、候補、共享資源、改期與套票／優惠碼原子驗收"],
        ["通知與 CRM", "通過", "同日提醒、失敗重試、跨渠道獨立、opt-in 與去重通過"],
        ["UI/UX", "已優化", "登入落點、權限提示、系統任務分區、品牌設定分區與手機版完成"],
        ["正式外部渠道", "待品牌資料", "需真實 LINE／LIFF、Email、綠界或藍新、DNS 證據"],
    ], [1.3, 1.0, 4.4])

    doc.add_heading("本輪已修正的流程", level=2)
    add_table(doc, ["原問題", "優化後", "營運價值"], [
        ["品牌登入落在預約清單", "直接進入今日工作台", "先處理今天的待辦與阻擋"],
        ["未授權頁面靜默跳走", "返回工作台／總覽並說明原因", "減少員工誤認為系統故障"],
        ["品牌設定全部塞在長頁", "拆成五個工作區", "降低尋找成本與誤改風險"],
        ["系統總覽混合建立與交付", "拆成總覽／建立／交付", "系統管理者依任務完成交接"],
        ["LINE 啟用被當成就緒", "檢查渠道驗證與 LIFF ID", "不會在渠道未完成時誤判可上線"],
        ["訂金文案說不串金流", "拆分訂金規則與標準金流", "符合綠界／藍新現行能力"],
    ], [2.0, 2.15, 2.55])

    doc.add_heading("角色模型（請維持這個版本）", level=2)
    add_table(doc, ["管理身分", "完整責任", "可新增的員工"], [
        ["系統管理者", "跨品牌租戶、平台人員、健康、報表、稽核、系統政策", "系統員工；依平台總覽、品牌管理、方案、健康、報表、稽核、設定授權"],
        ["品牌管理者", "該品牌全部設定、營運資料、渠道與品牌人員", "品牌員工；依品牌設定、日常營運、指派工作授權"],
    ], [1.25, 3.0, 2.45])
    add_callout(doc, "重要", "系統員工與品牌員工不是第三、第四種管理身分；它們是兩個管理層各自授權的工作帳號。", "warning")

    doc.add_heading("本輪功能交付狀態", level=2)
    add_table(doc, ["階段", "狀態", "已完成"], [
        ["UI/UX 與角色", "staging PASS", "四身分桌機／手機、導覽、登入落點、切換按鈕、未授權返回與提示"],
        ["Phase 1", "staging PASS", "CSV 匯入、權限範本、渠道測試、來源轉換、交班待辦篩選"],
        ["Phase 2", "staging PASS", "自訂表單／同意、重複預約、主服務＋多加購、訊息範本、快速再約"],
        ["三品牌觀察", "等待真實使用", "觀察名單、四 KPI 與付費意願紀錄已完成；目前不以 QA fixture 冒充樣本"],
        ["Phase 3 加購", "尚未決策", "至少兩個真實品牌 positive／won 才進入報價，否則暫緩"],
    ], [1.2, 1.2, 4.5])

    doc.add_heading("工程 Review 摘要", level=2)
    add_table(doc, ["優先級", "剩餘風險", "建議"], [
        ["P1", "app/admin/actions.ts 超過 2,000 行", "按 booking、event、customer、channel、settings 分檔，保留 action 契約"],
        ["P1", "app/book/page.tsx 超過 1,100 行", "抽出步驟元件與顯式狀態機，再補逐步 browser 測試"],
        ["P1", "系統總覽跨品牌資料未聚合分頁", "改 aggregate RPC／view、搜尋與分頁，避免規模成長後線性載入"],
        ["P1", "外部渠道尚無正式帳號證據", "每品牌保存測試時間、結果與修復指引"],
        ["P2", "角色瀏覽器驗收尚未固定在 CI", "四身分登入、導覽、拒絕與切換做 staging deploy gate"],
    ], [0.6, 2.65, 3.8])
    add_callout(doc, "依賴安全", "npm audit --omit=dev：0 個已知正式依賴漏洞。TypeScript strict type-check 與契約驗證通過。", "success")

    add_section_break(doc, "B. 系統管理者手冊", "系統管理者負責建立與治理品牌，不代替品牌做日常預約或顧客操作。")
    add_manual_screen(
        doc, "1. 共用登入入口", "01-login.png", "圖 1　共用登入頁：先選工作區，再輸入帳號密碼", "同一登入頁提供品牌與系統入口；選擇入口不會在前端授予權限，實際仍由 server 判定。",
        ["開啟 `/admin/login`。", "系統管理者或系統員工選「系統管理後台」；品牌人員選「品牌營運後台」。", "輸入 Email 與密碼後登入。", "若帳號沒有該工作區權限，系統會登出並顯示拒絕原因。"],
        "不要把測試密碼、service role、LINE access token 或金流密鑰寫進手冊、客服訊息或品牌資料庫。",
    )
    add_manual_screen(
        doc, "2. 系統營運總覽", "02-system-overview.png", "圖 2　系統總覽：任務分類、品牌開通流程與跨品牌指標", "系統管理者登入後先判斷平台健康與品牌交付，不需要捲過建立表單與所有品牌卡。",
        ["確認右上角身分顯示「系統管理者」。", "閱讀三步品牌開通流程。", "檢查品牌總數、開通完成率、累計預約／報名、服務數與投遞失敗。", "依工作需要進入系統人員、營運健康、跨品牌報表或系統稽核。"],
    )
    add_manual_screen(
        doc, "3. 建立新品牌", "03-system-create-brand.png", "圖 3　建立品牌工作區：租戶與品牌管理者一次建立", "品牌建立與日常品牌設定已分開；系統層只建立租戶、預設設定與第一位品牌管理者。",
        ["切換到「建立新品牌」。", "輸入品牌名稱與小寫英數／連字號品牌代號。", "輸入品牌管理者 Email；這個人會收到邀請並成為品牌管理者。", "電話、地址或備註可先留白。", "按「建立品牌並寄送登入邀請」。", "建立後前往品牌交付狀態，確認品牌管理者是否完成後續設定。"],
        "正式建立前先確認品牌名稱、slug 與管理者 Email；不要用相近品牌名稱猜測目標。",
    )
    add_manual_screen(
        doc, "4. 品牌交付狀態", "04-system-brand-delivery.png", "圖 4　品牌交付工作區：每個品牌的完成度與下一步", "品牌交付卡把系統責任與品牌責任分開，系統管理者只處理租戶狀態、方案與合約備註。",
        ["切換到「品牌交付狀態」。", "查看啟用狀態與待完成項目。", "若缺品牌管理者或品牌短網址，由系統層處理；服務、排班、入口與員工權限交給品牌管理者。", "需要暫停租戶時使用軟停用，不刪除歷史預約。", "加購旗標只記錄合作狀態，不代表功能已完成。"],
    )
    add_manual_screen(
        doc, "5. 系統人員與權限", "05-system-permissions.png", "圖 5　系統人員：兩種身分與細權限", "系統管理者擁有完整跨品牌權限；其他工作帳號統一是系統員工，再依責任授權。",
        ["輸入員工 Email。", "帳號身分選「系統員工」，除非該人確實需完整系統管理責任。", "只勾選工作需要的權限，例如營運健康或跨品牌報表。", "按「儲存人員權限」。", "定期從系統稽核檢查人員與品牌狀態異動。"],
        "不要以『方便』為由全部勾選；品牌日常人員應由品牌管理者建立，不應放進系統人員。",
    )
    add_manual_screen(
        doc, "6. 三品牌試用觀察", "25-trial-observation.png", "圖 6　試用觀察：最多三個品牌、四個 KPI 與付費意願", "系統管理者只能選入已同意參與的真實品牌；平台以資料庫原子限制最多三個，不使用前端計數假裝限制。",
        ["到系統報表開啟試用觀察。", "選入已開始試用的真實品牌並設定觀察期間。", "確認首次可預約時間、設定頁退出率、權限求助與預約完成率開始累積。", "訪談後為每個加購記錄 positive、neutral、negative 或 won 與理由。", "至少兩個品牌 positive／won 才進入報價；資料不足時維持等待。"],
        "QA Role Matrix Brand 與 staging 測試品牌只供驗收，不計入三品牌市場結論。",
    )
    add_manual_screen(
        doc, "7. 系統員工權限範本", "26-system-permission-presets.png", "圖 7　系統人員權限範本與可調整細權限", "以平台營運、品牌交付、報表與稽核等工作範本開始，再按實際責任微調。",
        ["輸入既有系統員工 Email。", "先選最接近工作的權限範本。", "檢查自動勾選項目，移除非必要權限。", "儲存後以員工帳號登入，直接測試允許與拒絕頁面。", "工作異動時立即收回舊權限。"],
    )

    add_section_break(doc, "C. 品牌管理者手冊", "品牌管理者用七步啟用中心完成上線，之後以今日工作台進行日常營運。")
    add_manual_screen(
        doc, "1. 今日工作台與七步啟用", "06-brand-dashboard.png", "圖 6　品牌管理者今日工作台：上線阻擋與今日任務", "品牌登入後直接看到營運任務；新品牌還會看到七步啟用中心。",
        ["先處理標示『阻擋』的啟用項目。", "服務／活動、人員／資源／排班完成後，才開放顧客入口。", "LINE 未正式驗證時可先用瀏覽器備援測試，不可宣稱 LINE 已上線。", "完成通知與付款後執行正式上線 smoke test。", "上線後每天先看待確認、待付款、通知失敗與未到。"],
    )
    add_manual_screen(
        doc, "2. 品牌資料", "07-brand-settings-profile.png", "圖 7　品牌設定：五個任務工作區與渠道狀態", "設定首頁只顯示品牌資料與跨領域就緒狀態，不再把全部表單堆在同一頁。",
        ["開啟「設定中心 → 品牌與規則」。", "在「品牌資料」填名稱、短網址、官方帳號 ID、destination、電話、地址與簡介。", "確認上方 LINE、金流、Email、網域狀態；『待處理』不等於故障，而是尚未完成正式條件。", "儲存後以公開頁檢查顧客實際看到的品牌資訊。"],
        "LINE Webhook destination 用於多品牌路由；請從實際 webhook payload 取得，不要自行猜值。",
    )
    add_manual_screen(
        doc, "3. 預約與入口規則", "08-brand-settings-booking.png", "圖 8　預約規則：標準模組、時間／場次模式與顧客規則", "所有品牌差異仍由設定驅動；時間制與場次制都保留，不在程式中寫死單一模式。",
        ["切換到「預約與入口規則」。", "決定是否啟用活動、會員套票與 CRM Lite。", "選擇時間制或場次制。", "設定首次服務延長、一電話多人、訂金、預約前置與取消／改期時限。", "最後才開啟公開預約或報名入口並儲存。"],
        "開啟訂金只建立收款規則；若需要線上付款，還必須在『付款與通知』完成綠界或藍新設定。",
    )
    add_manual_screen(
        doc, "4. 付款與通知", "09-brand-settings-channels.png", "圖 9　付款與通知：LINE、標準金流與 Email 分區", "三個外部渠道各自顯示狀態與必要條件；一個渠道失敗不應阻斷另一渠道。",
        ["先到 LINE 設定完成渠道驗證。", "需要線上訂金時選綠界或藍新、測試／正式環境與 Merchant ID。", "HashKey／HashIV 由部署環境提供，不會回填到畫面。", "Email 提醒需部署 Resend 金鑰與已驗證寄件網域。", "每個渠道完成後都要記錄一次實際測試結果。"],
    )
    add_manual_screen(
        doc, "5. 團隊與品牌員工權限", "10-brand-permissions.png", "圖 10　品牌團隊：品牌管理者與品牌員工細權限", "品牌管理者不可自行降級；其他帳號統一為品牌員工，權限以工作內容授予。",
        ["輸入員工 Email 與初始密碼。", "帳號身分選「品牌員工」。", "櫃檯勾「日常營運管理」；服務提供者勾「指派工作處理」並在資源頁完成指派；需要設定者才勾「品牌設定與整合」。", "按「新增帳號」。", "離職或調職時先移除品牌權限，不刪除歷史營運資料。"],
        "每個品牌至少保留一位品牌管理者；員工不能管理其他人員或自行提高權限。",
    )
    add_manual_screen(
        doc, "6. CSV 資料匯入", "41-csv-import-complete.png", "圖 11　CSV 顧客匯入完成：成功／失敗筆數與最近工作紀錄", "匯入精靈支援顧客、服務與套票餘額；先在瀏覽器預覽，再由 server 端按品牌規則寫入。",
        ["選擇顧客、服務或套票餘額。", "下載或展開欄位範例，準備 UTF-8 CSV。", "選檔後檢查前 10 筆預覽與總筆數。", "確認資料類型與目前品牌後再執行。", "查看成功、失敗與逐列原因；修正失敗列後用新工作重送。"],
        "同一匯入工作重送不會重複寫入；一電話多人上限仍由 server 端設定檢查。",
    )
    add_manual_screen(
        doc, "7. 渠道測試中心", "42-channel-test-results.png", "圖 12　渠道測試結果：LINE、LIFF、Email、金流與品牌入口分開判定", "渠道測試不會把『功能已啟用』當成『外部服務可用』，也不會讀出任何 secret。",
        ["按「執行全部測試」。", "先修正紅色失敗，再處理品牌確實要啟用的黃色待完成。", "使用卡片的「前往設定」回到來源頁。", "修正後重跑並保存新的時間與結果。", "最後仍須在手機、收件匣或測試付款頁完成真實外部 smoke test。"],
        "平台短網址會實際執行 host＋slug 品牌解析；僅顯示網址字串不算通過。",
    )
    add_manual_screen(
        doc, "8. 交班待辦與篩選", "31-handoff-tasks.png", "圖 13　交班待辦：狀態、到期與負責人篩選", "櫃檯把跨班次工作留在品牌範圍內，不用聊天訊息或紙條追蹤。",
        ["建立待辦標題、內容、優先級與到期時間。", "指定負責人；尚未決定時保留未指派。", "交班時使用未完成、逾期、今日到期與負責人篩選。", "處理後更新進行中或完成。", "敏感 PII 只放在既有顧客／預約資料，不複製到待辦標題。"],
    )
    add_manual_screen(
        doc, "9. 建立服務、表單與多服務加購", "32-service-form-addons.png", "圖 14　服務設定：預約自訂欄位與多服務加購", "主服務可加入文字、選項、日期與必填同意欄位，也可綁定多個延伸加購；加購分鐘納入容量計算。",
        ["建立主服務、時長、緩衝與預約目標。", "新增顧客欄位，為同意條款使用『同意條款（必填）』。", "調整欄位順序並儲存，讓預約保存文字快照。", "在多服務加購區選主服務，填名稱、增加分鐘與價格。", "從顧客入口驗證欄位、加購合計與可用時段。"],
        "目前交付的是一個主服務＋多個加購；跨多位服務人員的複雜組合仍需另立容量與付款設計。",
    )
    add_manual_screen(
        doc, "10. 顧客填寫自訂欄位與加購", "43-browser-custom-form-addons.png", "圖 15　瀏覽器顧客入口：自訂欄位、必填同意與多個加購選項", "顧客選服務後才看到該服務的欄位與加購；未完成必填同意或選項時，server 不會建立預約。",
        ["填姓名、電話、出生年月日與需要的 Email。", "選服務後完成預約前資料與必填同意。", "勾選需要的加購，確認增加分鐘與價格。", "選日期；系統依主服務＋加購總時長重新查詢時段。", "再選時段與是否連續週次。"],
    )
    add_manual_screen(
        doc, "11. 重複預約", "44-browser-recurring-booking.png", "圖 16　顧客已選時段後可建立 1–4 週重複預約", "品牌未啟用訂金時可開放重複預約；系統先確認每一週容量，再在同一交易中建立全部預約。",
        ["選擇可預約時段。", "從下拉選單選本次或連續週數。", "送出前再確認服務、加購與第一個日期。", "任一週沒有名額時整批失敗，不留下部分預約。", "成功後從我的預約逐筆取消或改期。"],
        "重複系列的每一筆仍有獨立預約狀態、容量、通知與稽核；訂金品牌目前請逐筆付款。",
    )
    add_manual_screen(
        doc, "12. 快速再次預約", "46-browser-quick-rebook.png", "圖 17　我的預約：三週系列與每筆『再次預約』入口", "快速再約帶回正確品牌與最近服務，但仍重新查詢最新排程與容量，不複製舊時段。",
        ["從完成頁點「查看我的預約」。", "在最近一次服務按「再次預約」。", "確認網址保留目前品牌且服務已預選。", "重新選日期、時段、欄位與加購。", "送出後回到我的預約確認新紀錄。"],
    )
    add_manual_screen(
        doc, "13. 訊息範本", "33-message-presets.png", "圖 18　非 AI 訊息範本：白名單變數與發送前預覽", "常用通知不需每次重打；範本只允許已知顧客／預約欄位，不執行任意模板程式。",
        ["從訊息中心選擇確認、提醒、改期或追蹤範本。", "確認可用變數與預覽。", "依本次情境調整文字。", "選擇收件人與渠道後再送出。", "檢查投遞結果，不將未同意顧客加入行銷訊息。"],
    )
    add_manual_screen(
        doc, "14. 品牌員工權限範本", "34-brand-permission-presets.png", "圖 19　品牌權限範本：櫃檯、服務提供者、行銷與財務檢視", "範本降低逐項勾選錯誤，但儲存前仍須按該員工實際工作微調。",
        ["先選櫃檯、服務提供者、行銷或財務檢視。", "核對勾選後移除不需要的設定／資料權限。", "服務提供者另在資源頁完成工作指派。", "儲存後用該員工帳號測試導覽與直接網址。", "職務變更時重新套用並收回舊權限。"],
    )
    add_manual_screen(
        doc, "15. 建立服務與方案（基礎欄位）", "11-brand-services.png", "圖 20　服務設定：跨產業預約目標與服務時長", "服務可以要求指定人員、可選人員／系統安排，或只占用場地設備；不要為資源型服務虛構人員。",
        ["開啟「服務與方案」。", "建立服務名稱、時長、緩衝與價格。", "選擇預約目標：必須指定、可指定／系統安排、或只使用資源。", "需要訂金或套票時設定適用範圍。", "儲存後再到人員與資源、服務排程建立可約容量。"],
    )
    add_manual_screen(
        doc, "16. 建立服務排程", "12-brand-schedules.png", "圖 21　服務排程：同日多段與資源型服務", "同一人同一天可有上午、下午、晚間多段；排程不可只保留第一段。",
        ["先確認服務提供者或資源已建立且啟用。", "選擇星期、開始／結束時間與服務。", "人員型服務選服務提供者；資源型服務可用 service_id 建立共用服務排程。", "同一天需要多段時分別新增。", "用例外日期處理休假、停班或臨時加開。", "最後從顧客入口驗證未來、未滿且符合前置時間的時段。"],
    )

    add_section_break(doc, "D. 日常營運手冊", "櫃檯與營運人員每天只處理待辦、顧客狀態與必要報表，不進入品牌安全設定。")
    add_manual_screen(
        doc, "1. 預約列表", "15-appointments.png", "圖 13　預約列表：日期、狀態與訂金", "時間制與場次制會依品牌設定顯示；列表保留取消歷史，不以刪除取代取消。",
        ["從今日工作台點待確認，或開啟「預約列表」。", "選日期與狀態篩選。", "確認顧客、服務、人員、初／複診、狀態與訂金。", "取消時使用狀態動作；改期使用正式改期流程，避免手動建立重複預約。", "遇到額滿或取消空缺時檢查預約候補。"],
    )
    add_manual_screen(
        doc, "2. 活動管理", "16-events.png", "圖 14　活動管理：場次、票種、表單與公開狀態", "活動報名與預約使用不同容量與候補狀態機，但共用品牌顧客與通知邊界。",
        ["建立活動名稱、slug、公開／私人入口與報名期間。", "建立一個或多個活動場次與容量。", "設定票種、價格、套票／優惠碼規則。", "建立自訂報名表單與條款同意。", "先在私人或測試狀態驗證，再發布公開入口。"],
    )
    add_manual_screen(
        doc, "3. 報名與 QR 報到", "17-registrations.png", "圖 15　報名管理：付款、候補、票券與報到狀態", "營運人員可處理待付款、候補與報到；同一 QR 重複掃描不會新增第二次報到。",
        ["依活動、場次、付款或報名狀態篩選。", "待付款逾時由排程取消並釋放優惠／候補名額。", "現場以 QR 報到；若 QR 錯誤可用受保護的人工搜尋。", "取消、未到與候補遞補都使用正式狀態動作。", "需要名單時使用授權匯出，避免在共用裝置保存 PII。"],
    )
    add_manual_screen(
        doc, "4. CRM Lite", "18-crm-lite.png", "圖 16　CRM Lite：分眾、互動時間軸與三種規則自動化", "CRM Lite 只做同意管理、規則分眾、互動紀錄與生日／完成後／久未到自動化，不擴成完整 CRM。",
        ["先確認顧客已同意 marketing opt-in。", "建立或預覽分眾條件。", "檢查顧客互動時間軸與封鎖狀態。", "啟用生日、服務完成後或久未到自動化。", "查看每渠道 sent／failed／skipped／duplicate 原因。", "顧客取消同意後停止行銷投遞。"],
    )
    add_manual_screen(
        doc, "5. 營運報表", "19-reports.png", "圖 17　品牌報表：預約、報名、付款、通知與漏斗", "報表依品牌與日期範圍統計；品牌員工永遠不能看到其他租戶資料。",
        ["選擇日期範圍與需要的營運維度。", "先看預約量、取消／未到率與服務／人員分佈。", "活動則看報名、候補、付款與報到。", "比較通知失敗與預約轉換，找出渠道問題。", "匯出前再次確認目前品牌與日期範圍。"],
    )

    add_section_break(doc, "E. LINE 與 Rich Menu 完整流程", "Rich Menu 是顧客任務入口，不是只有一張圖片；正式發布前必須先完成渠道、草稿、圖片、動作與驗證。")
    add_manual_screen(
        doc, "1. LINE／LIFF 渠道設定", "13-line-setup.png", "圖 18　LINE 渠道：啟用、模式、destination、Login Channel、LIFF 與外部驗證", "每品牌保存非機密識別資料；channel secret 與 access token 只存在 server environment。",
        ["開啟「LINE／LIFF 連線」。", "選共享渠道或品牌獨立渠道。", "填 Webhook destination、LINE Login Channel ID、LIFF ID 與 endpoint path。", "在部署環境設定對應的 channel secret 與 access token。", "按「重新驗證渠道」。", "只有外部驗證成功後，才進入 Rich Menu 發布。"],
        "LINE Official Account Manager 與 Messaging API 建立的同一個 Rich Menu 不能互相編修；XINHOW 發布的版本應持續由本平台管理。",
    )
    add_manual_screen(
        doc, "2. 建立 Rich Menu 草稿", "14-richmenu-workspace.png", "圖 19　Rich Menu 工作區：發布前檢查、模板、版型與線上版本", "草稿保存不會改動線上選單；發布前檢查會逐項標示 destination、token、Login Channel、LIFF 與渠道驗證。",
        ["輸入草稿名稱與選單列文字。", "依已啟用模組選快速模板。", "選擇格數與版型；完整六格建議圖片為 2500 × 1686 px。", "確認右側目前線上版本與版本紀錄。", "發布前先把所有橘色檢查項目處理完成。"],
    )
    add_manual_screen(
        doc, "3. 設定每格顧客動作", "14b-richmenu-actions.png", "圖 20　每格動作：名稱、無障礙標籤、任務與可選 URI", "格位應對應顧客要完成的任務，例如立即預約、我的預約、品牌資訊、票券、會員與客服。",
        ["為每格輸入顧客可理解的顯示名稱。", "填寫 20 字以內的無障礙標籤。", "從動作選單選標準顧客任務。", "只有特殊需求才填自訂 URI；一般任務使用系統產生的品牌安全入口。", "逐格檢查後按「另存草稿版本」。"],
        "不要把內部管理頁、未驗證外部網址或其他品牌 slug 放進 Rich Menu。",
    )
    add_manual_screen(
        doc, "4. 圖片、預覽、驗證與發布", "14c-richmenu-publish.png", "圖 21　Rich Menu 圖片與草稿發布區的實際畫面", "系統會裁切並壓縮圖片，但內容與文字可讀性仍由品牌負責；發布前先用瀏覽器逐格測試。",
        ["上傳符合版型比例的 PNG／JPEG。", "確認系統偵測的實際 MIME、尺寸與檔案大小。", "查看點擊區覆蓋預覽，確認文字沒有跨格。", "逐格按「瀏覽器測試」。", "按驗證；只有通過的草稿才可發布。", "發布成功後再用手機 LINE 重新開啟聊天室確認。"],
    )
    add_manual_screen(
        doc, "5. Alias 頁籤與顯示排程", "14d-richmenu-versioning.png", "圖 22　Rich Menu Alias 頁籤與 Asia/Taipei 顯示排程", "Alias 可在同一 LINE Channel 內切換多個已上傳版本；排程會切換指定版本並在結束後回復先前版本。",
        ["先發布並取得可用的 LINE Rich Menu ID。", "建立同 channel 唯一的 Alias ID 與管理名稱。", "選擇同一品牌、同一 channel 的版本後同步 Alias。", "需要期間活動時選顯示版本、台北開始與結束時間。", "查看排程成功／失敗與重試紀錄。", "成效檢視使用 LINE 官方 impressions／clicks 與平台匿名轉換，不推算被官方抑制的數字。", "若新版表現或內容有問題，從版本紀錄回復上一個正式版。"],
        "LINE 手機端重新開啟聊天室後才會看到新的預設 Rich Menu；更新可能需要短暫時間，不要用桌面版 LINE 當正式視覺驗收。",
    )

    add_section_break(doc, "F. 員工與手機操作", "導覽隱藏與 server 權限同時生效；員工遇到未授權功能時會知道下一步，不再誤以為系統壞掉。")
    add_manual_screen(
        doc, "1. 系統員工遇到未授權功能", "37-system-employee.png", "圖 23　系統員工未授權：返回系統總覽並顯示申請方式", "系統員工只取得被授權的平台工作；直接輸入系統人員頁仍會被 server guard 拒絕。",
        ["從系統總覽進入被授權的工作。", "若顯示權限提醒，確認工作是否屬於目前職責。", "確有需要時由系統管理者套用或調整權限範本。", "不要把系統管理者帳號共用給營運人員。"],
    )
    add_manual_screen(
        doc, "2. 品牌員工遇到未授權功能", "39-brand-employee.png", "圖 24　品牌員工未授權：返回工作台並顯示申請方式", "品牌員工只看得到被授權的日常工作；直接輸入設定網址也不會繞過 server guard。",
        ["員工先從今日工作台與左側導覽進入工作。", "若系統顯示權限提醒，確認該工作是否真的屬於此員工。", "確有需要時由品牌管理者到團隊與權限調整。", "不要共用品牌管理者帳號，也不要把管理者密碼交給員工。"],
    )
    add_mobile_pair(doc)
    add_four_role_mobile_grid(doc)

    add_section_break(doc, "G. 產品路線與新增功能", "下一步先降低導入與操作成本，再做成交與回訪；超出 v3 的大型能力另立商業決策。")
    doc.add_heading("市場對照結論", level=2)
    p = doc.add_paragraph()
    set_font(p.add_run("台灣同類產品多強調 LINE 自助預約、提醒、會員、訂金與簡單報表；國際成熟產品再加入資料匯入、自訂表單、重複／多服務預約、員工營運、no-show 政策、POS 與多渠道整合。XINHOW 的差異不應是『選單最多』，而應是："), size=10, color=SLATE)
    add_callout(doc, "建議定位", "以 LINE 為主要入口、同時支援瀏覽器備援的多品牌預約、活動報名與 CRM Lite 營運平台。", "info")
    doc.add_heading("Phase 1：已交付導入與日常工具", level=2)
    add_table(doc, ["功能", "staging 交付", "KPI"], [
        ["CSV 匯入精靈", "顧客、服務、套票；預覽、逐列結果、工作冪等與品牌限制", "首次可預約時間、匯入成功率"],
        ["員工權限範本", "系統與品牌兩層範本；套用後仍可微調，server guard 不變", "設定時間、權限求助與越權嘗試"],
        ["渠道測試中心", "LINE／LIFF／Email／金流／公開入口分項結果與修復連結", "渠道啟用成功率、修復時間"],
        ["來源與轉換總覽", "來源瀏覽、開始、完成、轉換率；品牌與日期範圍隔離", "每來源預約完成率"],
        ["交班待辦篩選", "負責人、到期、狀態、優先級與逾期篩選", "待處理平均時間"],
    ], [1.65, 3.75, 1.45])
    doc.add_heading("Phase 2：已交付成交與回訪工具", level=2)
    add_table(doc, ["功能", "staging 交付", "重要邊界"], [
        ["預約自訂表單與同意欄位", "文字、長文字、日期、選項、勾選、必填同意與規則快照", "不做病歷或 HIS"],
        ["重複預約", "顧客可選 1 至品牌上限週次；全部有容量才原子建立", "每筆獨立狀態與通知；任一失敗整批回復"],
        ["多服務／加購", "一個主服務可選多個加購；加購分鐘與價格納入檢查／快照", "跨多位人員的複雜組合尚未交付"],
        ["非 AI 訊息範本", "確認、提醒、改期與追蹤範本；白名單變數與預覽", "不擴成完整 CRM 或 AI 行銷"],
        ["快速再預約", "從最近服務／人員帶回品牌預約入口，再重新查可用時段", "沿用既有顧客身分與租戶驗證"],
    ], [1.7, 3.35, 1.8])
    doc.add_heading("Phase 3：另行報價與架構評估", level=2)
    add_bullets(doc, [
        "Google／Outlook 雙向行事曆同步。",
        "完整退款、對帳與財務結算。",
        "POS、商品庫存、員工分潤與薪資。",
        "多語系、進階白牌與原生管理 App。",
        "評價導流、推薦獎勵與公開活動市集。",
        "AI 行銷文案、AI 自動分眾與全自動投放。",
    ])
    add_callout(doc, "不要直接塞進現有導覽", "Phase 3 每一項都會改變商業模式、資料模型或權限，必須各自建立目標、範圍、價格與驗收標準。", "warning")

    doc.add_heading("三品牌付費意願決策規則", level=2)
    add_table(doc, ["資料狀態", "決策"], [
        ["少於三個真實品牌或訪談未完成", "等待；不得用 QA fixture 或工程判斷代替市場證據"],
        ["至少兩個品牌為 positive／won", "建議進入 discovery、報價與獨立驗收規格"],
        ["三品牌齊全但少於兩個 positive／won", "暫緩，不進開發排程"],
    ], [3.2, 3.5])

    doc.add_heading("建議執行順序", level=2)
    add_numbered(doc, [
        "UI/UX、四身分、Phase 1 與 Phase 2 已完成 staging 技術驗收。",
        "由系統管理者選入三個已同意的真實試用品牌並開始觀察。",
        "每週檢查首次可預約時間、設定退出率、權限求助與預約完成率，並記錄改善原因。",
        "完成三品牌訪談與加購意願後套用兩票門檻。",
        "只對達門檻的行事曆、退款對帳、POS／分潤、多語或白牌建立報價與獨立規格。",
    ])

    add_section_break(doc, "H. 驗收、上線與故障排除", "本地 build、staging 內部測試與正式外部渠道是三種不同證據；只有全部符合時才能宣稱完整上線。")
    doc.add_heading("2026-08-13 staging 驗收證據", level=2)
    add_table(doc, ["Gate", "結果", "證據摘要"], [
        ["TypeScript", "PASS", "tsc --noEmit"],
        ["契約測試", "PASS", "角色、租戶、Rich Menu、候補、付款、通知與 RLS 契約"],
        ["依賴安全", "PASS", "npm audit --omit=dev：0 vulnerabilities"],
        ["公開與 Cron 邊界", "PASS", "公開頁 200；未授權 Cron 401"],
        ["RLS／跨品牌", "PASS", "anon PII 拒絕、跨租戶讀寫拒絕、品牌設定最小權限"],
        ["預約／候補／資源", "PASS", "時間與場次併發、候補、共享資源、改期、訂金與 provider scope"],
        ["活動／付款／會員", "PASS", "私人入口、容量、QR、套票、優惠碼、逾時與 webhook 冪等"],
        ["通知／CRM", "PASS", "同日提醒、重試、去重、opt-in、三種規則與跨渠道獨立"],
        ["瀏覽器身分", "PASS", "本人範圍、跨品牌拒絕、host spoof 與 token 竄改拒絕"],
        ["四角色 UI", "PASS", "系統管理者、系統員工、品牌管理者、品牌員工實際 staging 推演"],
        ["Phase 1／2 後端", "PASS", "CSV 冪等、加購時長、三週系列、容量失敗整批回復與三品牌原子上限"],
        ["Phase 1／2 前端", "PASS", "CSV 完成、渠道結果、自訂表單、加購、重複預約與快速再約 staging 操作"],
    ], [1.5, 0.65, 4.6])
    add_callout(doc, "測試執行規則", "staging 完整 gate 與個別 notification gate 必須序列執行。若上一支仍在清理暫存資料就啟動下一支，可能形成 QA 清理競爭；這不是產品資料競爭，但會污染驗收結果。", "warning")

    doc.add_heading("正式上線前必要外部證據", level=2)
    add_table(doc, ["項目", "必要證據", "未完成時"], [
        ["LINE Messaging API", "真實 destination、access token、webhook 驗證、訊息實送", "維持 fail-closed；只用瀏覽器備援"],
        ["LINE Login／LIFF", "真實 Channel ID、LIFF ID、手機登入、品牌綁定", "不可宣稱 LINE 顧客身分流程完成"],
        ["Rich Menu", "圖片與逐格動作、手機顯示、發布／回復、Alias／排程（若使用）", "保留草稿，不發布"],
        ["Email", "Resend 金鑰、寄件網域驗證、實際顧客測試信", "Email 保持停用；LINE 不受影響"],
        ["綠界／藍新", "正式商店、密鑰、付款、回呼、重送冪等與逾時", "只記錄付款狀態，不開正式線上收款"],
        ["DNS／自訂網域", "TXT 驗證、HTTPS、host→品牌解析與 spoof 拒絕", "先使用平台短網址"],
    ], [1.3, 3.5, 2.0])

    doc.add_heading("管理者交付檢查單", level=2)
    add_table(doc, ["負責層", "交付前必勾"], [
        ["系統管理者", "□ 品牌／slug 正確　□ 第一位品牌管理者登入　□ 預設設定完成　□ 套餐／加購未誤標交付　□ 真實試用同意"],
        ["品牌管理者", "□ 服務／活動　□ 人員／資源　□ 多段排程　□ CSV　□ 員工最小權限　□ 表單／加購／重複規則"],
        ["渠道", "□ 渠道中心　□ LINE／LIFF 手機　□ Rich Menu 發布／回復　□ Email 收件　□ 測試付款　□ DNS／HTTPS"],
        ["Smoke test", "□ 桌機／手機預約　□ 改期／取消／候補　□ 通知　□ 我的預約／快速再約　□ 跨租戶拒絕　□ 報表／匯出"],
    ], [1.35, 5.35])
    add_callout(doc, "可編輯完整版", "同一交付包另附 admin-delivery-checklist.md，可逐品牌填寫 URL、測試時間、負責人、未完成項目與簽核日期。", "info")

    doc.add_heading("外部渠道設定順序", level=2)
    add_numbered(doc, [
        "LINE：確認 Messaging API 與 LINE Login Channel，從 webhook payload 取得 destination，再設定 Login Channel ID、LIFF ID 與 endpoint。",
        "Server environment：依共享或品牌獨立模式設定 secret／access token；不得保存於品牌資料表或文件。",
        "Rich Menu：草稿、圖片、逐格動作、瀏覽器測試、驗證、發布、手機顯示與回復依序完成。",
        "Email：驗證寄件網域與寄件人、設定 API key、開啟品牌渠道、實際收件與去重測試。",
        "綠界／藍新：先測試商店、付款、return、notify 重送冪等與逾時，再決定正式啟用。",
        "網址：平台 host 加入共享清單；自訂網域完成 DNS token、verified_at、HTTPS 與跨品牌 spoof 拒絕。",
        "回到渠道測試中心保存最後結果；伺服器通過後仍完成手機、收件與付款 smoke test。",
    ])
    add_callout(doc, "可編輯完整版", "同一交付包另附 external-channel-setup.md，包含 LINE、Rich Menu、Email、綠界／藍新、短網址、自訂網域、嵌入與狀態判讀。", "info")

    doc.add_heading("常見問題", level=2)
    qa = [
        ("登入後看不到設定中心", "品牌員工沒有品牌設定權限。若工作確實需要，請品牌管理者到團隊與權限調整；不要共用管理者帳號。"),
        ("LINE 已啟用但仍顯示待完成驗證", "啟用只是產品開關；仍需 destination、access token、Login Channel、LIFF ID 與外部驗證。"),
        ("Rich Menu 儲存後線上沒有變", "儲存只建立草稿。必須上傳圖片、逐格測試、通過驗證並發布。"),
        ("新 Rich Menu 手機沒有立刻出現", "重新開啟 LINE 聊天室；預設選單變更可能需要短暫時間。請勿用桌面版 LINE 當最終驗收。"),
        ("預約頁沒有時段", "依序檢查服務啟用、人員／資源、同日排程、例外日期、前置時間、最大預約天數、容量與公開入口。"),
        ("通知只有一個渠道失敗", "LINE 與 Email 獨立記錄與重試。先看 failed／skipped 原因，不要把另一渠道也關閉。"),
        ("金流啟用後仍顯示缺少密鑰", "Merchant ID 可在後台保存，但 HashKey／HashIV 必須由 server environment 提供。"),
        ("系統總控台與品牌後台怎麼切換", "同時具兩層權限的帳號，右上角才會顯示切換按鈕；返回品牌後台會進入今日工作台。"),
    ]
    for question, answer in qa:
        p = doc.add_paragraph()
        set_paragraph_keep(p)
        set_font(p.add_run(question), size=10.5, bold=True, color=NAVY)
        p = doc.add_paragraph()
        set_font(p.add_run(answer), size=9.8, color=SLATE)

    doc.add_heading("市場資料來源", level=2)
    add_bullets(doc, [
        "LINE Rich Menu 官方說明：https://developers.line.biz/en/docs/messaging-api/rich-menus-overview/",
        "Fresha 功能總覽：https://www.fresha.com/en-GB/for-business/features",
        "Square Appointments：https://squareup.com/us/en/appointments/pricing",
        "SimplyBook.me：https://simplybook.me/en/booking-system-features-and-integrations",
        "Acuity Scheduling：https://www.squarespace.com/scheduling",
        "Bookisha：https://www.bookisha.com/",
        "客立樂：https://www.qlieer.com/pricing",
        "LINE Biz-Solutions 美業應用：https://tw.linebiz.com/smb/industry-application/beauty/",
    ])

    add_callout(doc, "交付結論", "系統已從『功能集合』重整為『系統開通、品牌開通、日常營運、渠道成長』四條工作路徑。下一階段請以導入速度與營運 KPI 決定功能，不再以選單數量決定完整度。", "success")
    return doc


def main():
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = build_document()
    doc.core_properties.title = "XINHOW 多品牌預約與報名 SaaS 系統 Review 與圖文操作手冊"
    doc.core_properties.subject = "系統管理者、品牌管理者、員工權限、LINE Rich Menu 與產品優化路線"
    doc.core_properties.author = "XINHOW Product & Engineering"
    doc.core_properties.keywords = "XINHOW, 預約, 報名, SaaS, LINE, Rich Menu, 操作手冊"
    finalize_table_accessibility(doc)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
