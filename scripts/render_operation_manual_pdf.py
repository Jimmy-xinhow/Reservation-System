from __future__ import annotations

from io import BytesIO
from pathlib import Path
from xml.sax.saxutils import escape

from docx import Document
from docx.oxml.ns import qn
from docx.table import Table as DocxTable, _Cell
from docx.text.paragraph import Paragraph as DocxParagraph
from PIL import Image as PILImage
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.pdfbase.pdfmetrics import registerFont
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import (
    Image,
    KeepTogether,
    PageBreak,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)


ROOT = Path(__file__).resolve().parents[1]
DOCX_PATH = ROOT / "docs" / "deliverables" / "xinhowsaas-system-review-and-operation-manual-2026-08-13.docx"
PDF_PATH = ROOT / "docs" / "deliverables" / "xinhowsaas-system-review-and-operation-manual-2026-08-13.pdf"

NAVY = colors.HexColor("#08263A")
BLUE = colors.HexColor("#2178C7")
SLATE = colors.HexColor("#475569")
LIGHT_BLUE = colors.HexColor("#F4F8FC")
PALE_BLUE = colors.HexColor("#EAF3FB")
WHITE = colors.white
FONT_REGULAR = "MicrosoftJhengHei"
FONT_BOLD = "MicrosoftJhengHeiBold"


def register_fonts() -> tuple[str, str]:
    registerFont(TTFont(FONT_REGULAR, r"C:\Windows\Fonts\msjh.ttc", subfontIndex=0))
    registerFont(TTFont(FONT_BOLD, r"C:\Windows\Fonts\msjhbd.ttc", subfontIndex=0))
    return FONT_REGULAR, FONT_BOLD


def build_styles(font_regular: str, font_bold: str):
    styles = getSampleStyleSheet()
    return {
        "body": ParagraphStyle(
            "BodyZH",
            parent=styles["BodyText"],
            fontName=font_regular,
            fontSize=9.3,
            leading=13.2,
            textColor=SLATE,
            spaceAfter=5,
            allowWidows=0,
            allowOrphans=0,
        ),
        "title": ParagraphStyle(
            "TitleZH",
            parent=styles["Title"],
            fontName=font_bold,
            fontSize=24,
            leading=31,
            textColor=NAVY,
            spaceAfter=16,
        ),
        "h1": ParagraphStyle(
            "Heading1ZH",
            parent=styles["Heading1"],
            fontName=font_bold,
            fontSize=18,
            leading=24,
            textColor=NAVY,
            spaceBefore=12,
            spaceAfter=8,
            keepWithNext=1,
        ),
        "h2": ParagraphStyle(
            "Heading2ZH",
            parent=styles["Heading2"],
            fontName=font_bold,
            fontSize=13.5,
            leading=18,
            textColor=BLUE,
            spaceBefore=10,
            spaceAfter=6,
            keepWithNext=1,
        ),
        "h3": ParagraphStyle(
            "Heading3ZH",
            parent=styles["Heading3"],
            fontName=font_bold,
            fontSize=11,
            leading=15,
            textColor=NAVY,
            spaceBefore=7,
            spaceAfter=4,
            keepWithNext=1,
        ),
        "caption": ParagraphStyle(
            "CaptionZH",
            parent=styles["BodyText"],
            fontName=font_regular,
            fontSize=8,
            leading=11,
            alignment=TA_CENTER,
            textColor=SLATE,
            spaceAfter=7,
        ),
        "bullet": ParagraphStyle(
            "BulletZH",
            parent=styles["BodyText"],
            fontName=font_regular,
            fontSize=9.2,
            leading=13,
            leftIndent=13,
            firstLineIndent=-8,
            textColor=SLATE,
            spaceAfter=3,
        ),
        "table": ParagraphStyle(
            "TableZH",
            parent=styles["BodyText"],
            fontName=font_regular,
            fontSize=7.3,
            leading=10.3,
            textColor=SLATE,
        ),
        "table_header": ParagraphStyle(
            "TableHeaderZH",
            parent=styles["BodyText"],
            fontName=font_bold,
            fontSize=7.3,
            leading=10.3,
            textColor=WHITE,
        ),
        "callout": ParagraphStyle(
            "CalloutZH",
            parent=styles["BodyText"],
            fontName=font_regular,
            fontSize=8.8,
            leading=12.5,
            textColor=NAVY,
        ),
    }


def iter_block_items(parent):
    parent_element = parent.element.body if isinstance(parent, Document().__class__) else parent._tc
    for child in parent_element.iterchildren():
        if child.tag == qn("w:p"):
            yield DocxParagraph(child, parent)
        elif child.tag == qn("w:tbl"):
            yield DocxTable(child, parent)


def paragraph_has_page_break(paragraph: DocxParagraph) -> bool:
    return bool(paragraph._p.xpath(".//w:br[@w:type='page']"))


def paragraph_images(paragraph: DocxParagraph, document: Document):
    for blip in paragraph._p.xpath(".//a:blip[@r:embed]"):
        relationship_id = blip.get(qn("r:embed"))
        if relationship_id:
            yield document.part.related_parts[relationship_id].blob


def image_flowable(blob: bytes, max_width: float, max_height: float) -> Image:
    with PILImage.open(BytesIO(blob)) as source:
        width, height = source.size
    scale = min(max_width / width, max_height / height)
    return Image(BytesIO(blob), width=width * scale, height=height * scale)


def text_paragraph(text: str, style, prefix: str = "") -> Paragraph:
    safe = escape(text).replace("\n", "<br/>")
    return Paragraph(f"{prefix}{safe}", style)


def paragraph_flowables(paragraph: DocxParagraph, document: Document, styles: dict, usable_width: float):
    result = []
    image_blobs = list(paragraph_images(paragraph, document))
    for blob in image_blobs:
        result.append(image_flowable(blob, usable_width, 6.8 * inch))

    text = paragraph.text.strip()
    if text:
        style_name = paragraph.style.name if paragraph.style else "Normal"
        if style_name == "Title":
            style = styles["title"]
            prefix = ""
        elif style_name == "Heading 1":
            style = styles["h1"]
            prefix = ""
        elif style_name == "Heading 2":
            style = styles["h2"]
            prefix = ""
        elif style_name == "Heading 3":
            style = styles["h3"]
            prefix = ""
        elif style_name == "Caption":
            style = styles["caption"]
            prefix = ""
        elif style_name.startswith("List Bullet"):
            style = styles["bullet"]
            prefix = "• "
        elif style_name.startswith("List Number"):
            style = styles["bullet"]
            prefix = "• "
        else:
            style = styles["body"]
            prefix = ""
        result.append(text_paragraph(text, style, prefix))

    if paragraph_has_page_break(paragraph):
        result.append(PageBreak())
    elif not result:
        result.append(Spacer(1, 3))
    return result


def cell_flowables(cell: _Cell, document: Document, styles: dict, max_width: float):
    content = []
    for paragraph in cell.paragraphs:
        for blob in paragraph_images(paragraph, document):
            content.append(image_flowable(blob, max_width, 5.6 * inch))
        if paragraph.text.strip():
            content.append(text_paragraph(paragraph.text.strip(), styles["table"]))
    return content or [Spacer(1, 1)]


def table_flowable(table: DocxTable, document: Document, styles: dict, usable_width: float):
    row_count = len(table.rows)
    column_count = len(table.columns)
    if not row_count or not column_count:
        return Spacer(1, 1)

    column_width = usable_width / column_count
    data = []
    for row_index, row in enumerate(table.rows):
        rendered_row = []
        for cell in row.cells:
            if column_count == 1:
                rendered_row.append(cell_flowables(cell, document, styles, column_width - 12))
            elif any(paragraph_images(p, document) for p in cell.paragraphs):
                rendered_row.append(cell_flowables(cell, document, styles, column_width - 10))
            else:
                cell_style = styles["table_header"] if row_index == 0 else styles["table"]
                rendered_row.append(text_paragraph(cell.text.strip(), cell_style))
        data.append(rendered_row)

    rendered = Table(data, colWidths=[column_width] * column_count, repeatRows=1 if column_count > 1 else 0)
    commands = [
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("LEFTPADDING", (0, 0), (-1, -1), 6),
        ("RIGHTPADDING", (0, 0), (-1, -1), 6),
        ("TOPPADDING", (0, 0), (-1, -1), 6),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
        ("GRID", (0, 0), (-1, -1), 0.35, colors.HexColor("#CBD5E1")),
    ]
    if column_count == 1:
        commands.extend([
            ("BACKGROUND", (0, 0), (-1, -1), PALE_BLUE),
            ("BOX", (0, 0), (-1, -1), 0.8, BLUE),
        ])
    else:
        commands.append(("BACKGROUND", (0, 0), (-1, 0), NAVY))
        for row_index in range(1, row_count):
            if row_index % 2 == 0:
                commands.append(("BACKGROUND", (0, row_index), (-1, row_index), LIGHT_BLUE))
    rendered.setStyle(TableStyle(commands))
    return rendered


def draw_page(canvas, document):
    canvas.saveState()
    canvas.setFont(FONT_REGULAR, 7.5)
    canvas.setFillColor(NAVY)
    canvas.drawString(42, A4[1] - 27, "XINHOW | SYSTEM REVIEW & OPERATIONS GUIDE")
    canvas.setFillColor(SLATE)
    canvas.drawRightString(A4[0] - 42, 24, f"第 {document.page} 頁")
    canvas.setStrokeColor(colors.HexColor("#CBD5E1"))
    canvas.setLineWidth(0.4)
    canvas.line(42, A4[1] - 33, A4[0] - 42, A4[1] - 33)
    canvas.restoreState()


def build_pdf():
    document = Document(DOCX_PATH)
    font_regular, font_bold = register_fonts()
    styles = build_styles(font_regular, font_bold)
    pdf = SimpleDocTemplate(
        str(PDF_PATH),
        pagesize=A4,
        leftMargin=42,
        rightMargin=42,
        topMargin=43,
        bottomMargin=38,
        title="XINHOW 預約與報名 SaaS 系統 Review 與操作手冊",
        author="XINHOW Product & Engineering",
    )

    story = []
    for block in iter_block_items(document):
        if isinstance(block, DocxParagraph):
            flowables = paragraph_flowables(block, document, styles, pdf.width)
            if len(flowables) == 2 and isinstance(flowables[0], Image) and isinstance(flowables[1], Paragraph):
                story.append(KeepTogether(flowables))
            else:
                story.extend(flowables)
        else:
            story.append(table_flowable(block, document, styles, pdf.width))
            story.append(Spacer(1, 6))

    pdf.build(story, onFirstPage=draw_page, onLaterPages=draw_page)
    print(PDF_PATH)


if __name__ == "__main__":
    build_pdf()
